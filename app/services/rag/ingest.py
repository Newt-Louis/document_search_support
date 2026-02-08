import logging, sys, os, shutil, magic, csv, openpyxl, json
from typing import Generator, List, AsyncGenerator
from pathlib import Path
from fastapi import UploadFile, HTTPException

from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, StorageContext, Settings, Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.ingestion import IngestionPipeline
from llama_index.vector_stores.qdrant import QdrantVectorStore

from app.core.config import get_config

cfg = get_config()
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(Path(cfg.LOG_DIR)/"ingest.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)
CACHE_DIR = cfg.CACHE_DIR
ALLOWED_EXTENSIONS = {
    "application/pdf": ".pdf",
    "text/plain": ".txt",
    "text/csv": ".csv",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/msword": ".doc",
    "application/vnd.ms-excel": ".xls"
}

class IngestService:
    def __init__(self, vector_store: QdrantVectorStore, upload_dir: str):
        self.vector_store = vector_store
        self.upload_dir = upload_dir
        self.text_splitter = SentenceSplitter(chunk_size=1024, chunk_overlap=200)

    @staticmethod
    async def _validate_file(file:UploadFile):
        header = await file.read(2048)
        mime_type = magic.from_buffer(header, mime=True)
        logger.info(f">>> Phát hiện file có MIME type: {mime_type}")
        await file.seek(0)
        if mime_type not in ALLOWED_EXTENSIONS:
            if not (mime_type.startswith("text/") and file.filename.endswith((".txt", ".csv"))):
                logger.warning(f"File bị reject: {file.filename} (MIME: {mime_type})")
                raise HTTPException(
                    status_code=400,
                    detail=f"File giả mạo hoặc không hỗ trợ! Phát hiện định dạng thực tế: {mime_type}"
                )

    async def process_upload(self, file: UploadFile) -> AsyncGenerator[str, None]:
        try:
            await self._validate_file(file)
            yield json.dumps({"status": "validating", "progress": 5, "message": "Kiểm tra định dạng file xong."}) + "\n"
        except ValueError as e:
            yield json.dumps({"status": "error", "message": str(e)}) + "\n"
            raise HTTPException(status_code=400, detail=str(e))

        os.makedirs(self.upload_dir, exist_ok=True)
        file_path = os.path.join(self.upload_dir, file.filename)

        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            yield json.dumps({"status": "saving", "progress": 10, "message": "Đã lưu file lên server. Bắt đầu xử lý..."}) + "\n"
            async for progress_update in self.index_file(file_path):
                yield progress_update
            # success = await self.index_file(file_path)
            # if success:
            #     return f"Lập chỉ mục cho tài liệu {file.filename} thành công!"
            # return "Có lỗi khác exception không bắt được !!!"

        except Exception as e:
            logger.exception(f"Lỗi khi xử lý file {file.filename}")
            if os.path.exists(file_path):
                os.remove(file_path)
            yield json.dumps({"status": "error", "message": f"Lỗi hệ thống: {str(e)}"}) + "\n"

    async def index_file(self, file_path: str)->AsyncGenerator[str, None]:
        """
        Hàm index_file Generator xử lý index file lưu vào qdrant theo từng batch chạy theo tiến trình.
        """
        try:
            pipeline = IngestionPipeline(
                transformations=[
                    self.text_splitter,  # Cắt nhỏ (Chunking)
                    Settings.embed_model,  # Hóa vector (Embedding)
                ],
                vector_store=self.vector_store,  # Đẩy vào Qdrant
            )
            total_size = os.path.getsize(file_path)
            processed_size = 0
            yield json.dumps({"status": "processing", "progress": 15, "message": "Đang khởi tạo Pipeline..."}) + "\n"
            for doc_batch, file_progress in self._lazy_load_file(file_path,chunk_size_mb=5):
                if not doc_batch:
                    continue
                nodes = await pipeline.arun(documents=doc_batch, show_progress=False)
                ui_percent = 15 + int((file_progress / 100) * 80)
                if ui_percent > 95: ui_percent = 95
                yield json.dumps({
                    "status": "processing",
                    "progress": ui_percent,
                    "message": f"Đã vector hóa {len(nodes)} chunks dữ liệu..."
                }) + "\n"

                del doc_batch
                del nodes

            yield json.dumps({"status": "complete", "progress": 100, "message": "✅ Hoàn tất! Tài liệu đã sẵn sàng."}) + "\n"
        except Exception as e:
            logger.exception(f"Lỗi khi lập chỉ mục: {e}")
            if os.path.exists(file_path):
                os.remove(file_path)
            yield json.dumps({"status": "error", "message": f"Lỗi xử lý AI: {str(e)}"}) + "\n"
            raise HTTPException(status_code=500, detail=f"Lỗi khi Indexing: {str(e)}")

    @staticmethod
    def _lazy_load_file(file_path: str, chunk_size_mb: int = 10) -> Generator[tuple[List[Document], float], None, None]:
        """
        Đọc file theo từng phần nhỏ (Batch).
        chunk_size_mb: Kích thước mỗi lần đọc (Mặc định 10MB text)
        """
        ext = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path)

        if ext == ".txt":
            # Đọc file Text dòng theo dòng
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                batch_text = ""
                current_size = 0
                while True:
                    line = f.readline()
                    if not line: break

                    batch_text += line
                    current_size += len(line.encode('utf-8'))

                    if current_size >= chunk_size_mb * 1024 * 1024:
                        # Tính % dựa trên vị trí con trỏ file (Bytes read / Total bytes)
                        progress = min((f.tell() / file_size) * 100, 99)
                        yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path)})], progress
                        batch_text = ""
                        current_size = 0

                if batch_text:
                    yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path)})], 100

        elif ext == ".csv":
            # Đọc CSV dòng theo dòng
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                batch_text = ""
                current_size = 0
                for row in reader:
                    row_text = ", ".join(row) + "\n"
                    batch_text += row_text
                    current_size += len(row_text.encode('utf-8'))

                    if current_size >= chunk_size_mb * 1024 * 1024:
                        progress = min((f.tell() / file_size) * 100, 99)
                        yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path)})], progress
                        batch_text = ""
                        current_size = 0
                if batch_text:
                    yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path)})], 100

        elif ext == ".pdf":
            # Với PDF, ta dùng pypdf đọc từng trang
            import pypdf
            reader = pypdf.PdfReader(file_path)
            total_pages = len(reader.pages)
            if total_pages == 0: total_pages = 1
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    # Mỗi trang là 1 yield. Progress = trang hiện tại / tổng số trang
                    progress = ((i + 1) / total_pages) * 100
                    yield [Document(text=text,metadata={"page_label": str(i), "filename": os.path.basename(file_path)})], progress

        elif ext == ".xlsx":
            # Quan trọng: read_only=True giúp openpyxl không load hết vào RAM
            # data_only=True để lấy giá trị cuối cùng, không lấy công thức hàm
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            total_sheets = len(wb.sheetnames)

            for s_idx, sheet in enumerate(wb):
                batch_text = ""
                current_size = 0
                # Duyệt từng dòng trong sheet. Đây là Generator.
                for row in sheet.iter_rows(values_only=True):
                    # Lọc None và chuyển thành string, cách nhau bởi dấu phẩy
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = ", ".join(row_values) + "\n"

                    batch_text += row_text
                    current_size += len(row_text.encode('utf-8'))

                    # Kiểm tra ngưỡng chunk_size_mb
                    if current_size >= chunk_size_mb * 1024 * 1024:
                        progress = ((s_idx) / total_sheets) * 100 + 5  # +5 để nó nhích
                        yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path), "sheet": sheet.title})], progress
                        batch_text = ""  # Giải phóng RAM
                        current_size = 0

                # Yield phần còn dư của sheet hiện tại
                if batch_text:
                    progress = ((s_idx + 1) / total_sheets) * 100
                    yield [Document(text=batch_text,metadata={"filename": os.path.basename(file_path), "sheet": sheet.title})], progress
            wb.close()

        elif ext == ".docx":
            from docx import Document as DocxDocument
            # Word file là tập hợp các đoạn văn (Paragraphs)
            doc = DocxDocument(file_path)
            total_paras = len(doc.paragraphs)  # Lấy tổng số đoạn văn
            if total_paras == 0: total_paras = 1
            batch_text = ""
            current_size = 0

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text: continue

                text += "\n"
                batch_text += text
                current_size += len(text.encode('utf-8'))

                # Kiểm tra ngưỡng chunk_size_mb
                if current_size >= chunk_size_mb * 1024 * 1024:
                    progress = ((i + 1) / total_paras) * 100
                    yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path)})], progress
                    batch_text = ""
                    current_size = 0

            # Yield nốt phần còn dư
            if batch_text:
                yield [Document(text=batch_text, metadata={"filename": os.path.basename(file_path)})], 100
        else:
            yield []